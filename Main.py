# coding:utf-8
import sys
import docx
import shutil
import os
file=docx.Document('DOCX_PATH')
count=0
for i in range(len(file.paragraphs)):
 front=file.paragraphs[i].text.find('FRONT_KEYWORD')
 ending=file.paragraphs[i].text.find('END_KEYWORD')
 x=len(front)
 y=len(ending)
 filename=file.paragraphs[i].text[front+front:ending+ending]
 dst="OUTPUT_DST"+filename
 src="SOURCE_DST"+filename
 if(os.path.exists(src)):
  print(src)
  shutil.copyfile(src, dst)
  count=count+1
print('Output Done!')
print(count)
print(' files have been moved')
raw_input()

# Remove: os.remove(FILE_NAME)
